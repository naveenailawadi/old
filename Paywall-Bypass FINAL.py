from bs4 import BeautifulSoup as bs
import requests
from docx import Document


def paywall_bypass():
    # get the site
    website = str(input('What article do you want to read? \
Paste the link below. \n'))

    # name the site
    name = str(input('What do you want to call it?\n'))

    # create soup
    print('creating soup...')
    print('Do not touch anything.')
    site_raw = requests.get(website)
    site_soup = bs(site_raw.text, "html.parser")

    # write document
    print('writing document... \n')
    doc = Document()
    doc.save(name + '.docx')
    for header in site_soup.find_all('h1'):
        file = doc.add_paragraph("")
        file.add_run(header.text).bold = True
        doc.add_paragraph("\n")
    for p in site_soup.find_all('p'):
        doc.add_paragraph(p.text)

    # finish
    doc.save(name + '.docx')
    print('Your file has been created. \n It is called ' + name + '.doc')
    print('It will be found in finder under All My Files. \n')


# RUN IT
paywall_bypass()

# rerun?
answer = str(input('If you want to read another article, type "yes" and hit enter, \
Otherwise, type "no" and hit enter. \n'))
if 'yes' in answer.lower():
    paywall_bypass()
else:
    print('Enjoy!')
