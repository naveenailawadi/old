'''
This script keeps a reader briefed on the most recent news.
My selected sources were the Washington Post, New York Times, and Economist.
This script will pull the headlines from each of these sources
and then obtain the articles that are selected by the user.
Some of the lines have been commented out after I
switched the program to utilize multithreading.
'''

from bs4 import BeautifulSoup as bs
import requests
from docx import Document
import threading
import math


def print_grab_notifier(homepage):
    print('Grabbing headlines from ' + homepage + ' ... \n\n')
    return


def paywall_bypass(website):
    # create soup
    print('creating soup for ' + str(website))
    site_raw = requests.get(website)
    site_soup = bs(site_raw.text, "html.parser")

    # name document
    name_tag = site_soup.find('h1')
    name = name_tag.text

    # write document
    print('writing document... \n')
    doc = Document()
    doc.save(name + '.docx')
    for header in site_soup.find_all('h1'):
        file = doc.add_paragraph("")
        file.add_run(header.text).bold = True
        doc.add_paragraph("\n")
    for p in site_soup.find_all('p'):
        doc.add_paragraph(p.text)

    # finish
    doc.save(name + '.docx')
    print('Your file has been created. \n It is called ' + name + '.doc')
    print('It will be found in finder under All My Files. \n')


# dictionary of headlines to fill
headlines = {}

# these functions find headlines from a given news outlet
# and append themn to the master list


def headline_grab_washpost(lock):
    print_grab_notifier(homepages[0])
    website = homepages[0]
    # website = home
    # find websites
    site_raw = requests.get(website)
    site_soup = bs(site_raw.text, "html.parser")

    # finds headlines
    washpost_headline_top = site_soup.find('h1')
    washpost_headlines = site_soup.find_all('h2')

    # add headlines to master dict
    headlines[washpost_headline_top.text] = href_extractor(washpost_headline_top)
    for headline in washpost_headlines:
        headlines[headline.text] = href_extractor(headline.find('a'))
    for _ in range(100000):
        lock.acquire()
        increment()
        lock.release()
    return


def headline_grab_nyt(lock):
    print_grab_notifier(homepages[1])
    website = homepages[1]
    # website = home
    # find websites
    site_raw = requests.get(website)
    site_soup = bs(site_raw.text, "html.parser")

    # finds headlines
    nyt_headlines_raw = site_soup.find_all('a')
    nyt_headlines_filter = site_soup.find_all('h2')
    for headline in nyt_headlines_raw:
        headline_possible = headline.find('h2')
        # add headlines to master dict
        if headline_possible in nyt_headlines_filter:
            headlines[headline_possible.text] = str('https://www.nytimes.com/' + href_extractor(headline))
    for _ in range(100000):
        lock.acquire()
        increment()
        lock.release()
    return


def headline_grab_econ(lock):
    print_grab_notifier(homepages[2])
    website = homepages[2]
    # website = home
    # find websites
    site_raw = requests.get(website)
    site_soup = bs(site_raw.text, "html.parser")

    # finds headlines
    econ_headlines_raw = site_soup.find_all('a')
    econ_headlines_filter = site_soup.find_all('h3')
    for headline in econ_headlines_raw:
        headline_possible = headline.find('h3')
        # add headlines to master dict
        if headline_possible in econ_headlines_filter:
            headlines[headline_possible.text] = str('https://www.economist.com' + href_extractor(headline))
    for _ in range(100000):
        lock.acquire()
        increment()
        lock.release()
    return


def href_extractor(tag):
    tag_string = str(tag)
    start = tag_string.index('href=')
    after_start = int(start + 6)
    stop = tag_string.index('"', after_start)
    href = tag_string[start + 6: stop]
    return href


# news outlet URLs to scrape
homepages = ['https://www.washingtonpost.com/', 'https://www.nytimes.com/', 'https://www.economist.com/']

# printing function to let user know what is happening

'''
# obtain headlines from all websites

print_grab_notifier(homepages[0])
headline_grab_washpost(homepages[0])

print_grab_notifier(homepages[1])
headline_grab_nyt(homepages[1])

print_grab_notifier(homepages[2])
headline_grab_econ(homepages[2])
'''


def increment():
    """
    function to increment global variable x
    """
    global x
    x += 1


if __name__ == "__main__":
    global x
    # setting global variable x as 0
    x = 0

    lock = threading.Lock()

    # creating thread
    t1 = threading.Thread(target=headline_grab_washpost, args=(lock,))
    t2 = threading.Thread(target=headline_grab_nyt, args=(lock,))
    t3 = threading.Thread(target=headline_grab_econ, args=(lock,))

    # starting thread 1
    t1.start()
    # starting thread 2
    t2.start()
    # starting thread 3
    t3.start()

    # wait until thread 1 is completely executed
    t1.join()
    # wait until thread 2 is completely executed
    t2.join()
    # wait until thread 3 is completely executed
    t3.join()

    # both threads completely executed
    print("Headlines found! \n")


# print and index the headlines
print("Headlines:")
count = 0
for key, value in headlines.items():
    print(str(count) + ' : ' + key + '\n')
    count += 1

selected_headlines = []

# select articles to read
while True:
    article_choice = input('Copy one headline that your would like to read. \n\
If you are done, type "I am done". \n')
    if "i am done" in article_choice.lower():
        break
    # scan to check errors
    elif headlines[article_choice] not in selected_headlines:
        selected_headlines.append(headlines[article_choice])
    elif headlines[article_choice] in selected_headlines:
        print('You have already selected that article.')
    else:
        print('Input Error. Try again. \n\n')


# create word files of the websites
for selection in selected_headlines:
    paywall_bypass(selection)

print("Done!")
