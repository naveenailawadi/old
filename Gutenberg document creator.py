from bs4 import BeautifulSoup as bs
import requests
from docx import Document


def doc_creator():
    # get the site
    website = str(input('Paste the link of your online novel below. \n'))

    # create soup
    print('creating soup...')
    print('Do not touch anything.')
    site_raw = requests.get(website)
    site_soup = bs(site_raw.text, 'html.parser')

    # name the site
    title = site_soup.find('h1')
    name = title.text

    # write document
    print('writing document... \n')
    doc = Document()
    doc.save(name + '.docx')

    # find text
    chapters = site_soup.find_all('h2')

    # write document
    for chapter in chapters:
        file = doc.add_paragraph("")
        file.add_run(chapter.text).bold = True
        while '<p>' in str(chapter.find_next()):
            doc.add_paragraph(chapter.find_next().text)
            chapter = chapter.find_next()

    # finish
    doc.save(name + '.docx')
    print('Your file has been created. \n It is called ' + name + '.txt')
    print('It will be found in finder under All My Files. \n')
    return


doc_creator()
